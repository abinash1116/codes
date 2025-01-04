from docx import Document

def merge_word_files(base_file, additional_file, output_file):
    """
    Merges the content of `additional_file` into `base_file` and saves as `output_file`.
    """
    try:
        # Open the base document
        base_doc = Document(base_file)
        # Open the additional document
        additional_doc = Document(additional_file)
        
        # Append all paragraphs from the additional document to the base document
        for paragraph in additional_doc.paragraphs:
            base_doc.add_paragraph(paragraph.text)
        
        # Save the merged document
        base_doc.save(output_file)
        print(f"Successfully merged {additional_file} into {base_file}.")
        print(f"Merged document saved as: {output_file}")
    except Exception as e:
        print(f"An error occurred: {e}")

# Files to merge
base_file = "spotify.docx"          # The file to which content is added
additional_file = "images.docx"     # The file to be added
output_file = "merged_spotify.docx" # Output file name

# Call the function
merge_word_files(base_file, additional_file, output_file)
